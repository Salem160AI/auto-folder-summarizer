
import os
import collections
import re
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from tkinter import Tk, filedialog, simpledialog, messagebox
from nltk.corpus import stopwords
import nltk

nltk.download('stopwords')
stop_words = set(stopwords.words('english'))

def classify_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext in ['.pdf', '.docx', '.txt', '.jpg', '.jpeg', '.png']:
        return ext[1:]
    else:
        return 'other'

def extract_text(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    try:
        if ext == '.pdf':
            reader = PdfReader(file_path)
            return "\n".join([page.extract_text() or '' for page in reader.pages])
        elif ext == '.docx':
            doc = DocxDocument(file_path)
            return "\n".join([p.text for p in doc.paragraphs])
        elif ext == '.txt':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
    return ''

def count_words(text):
    words = re.findall(r'\b\w+\b', text.lower())
    return len(words)

def extract_keywords(text, top_n=5):
    words = re.findall(r'\b\w+\b', text.lower())
    filtered = [w for w in words if w not in stop_words and len(w) > 2]
    counter = collections.Counter(filtered)
    return [word for word, _ in counter.most_common(top_n)]

def summarize_folder(folder_path):
    file_data = []
    summary = collections.defaultdict(int)
    total_size = 0
    all_text = ""

    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if not os.path.isfile(file_path):
            continue
        ext = classify_file(file_path)
        size = os.path.getsize(file_path)
        total_size += size
        summary[ext] += 1

        if ext in ['pdf', 'docx', 'txt']:
            text = extract_text(file_path)
            word_count = count_words(text)
            keywords = extract_keywords(text)
            all_text += " " + text
        else:
            word_count = 0
            keywords = []

        file_data.append({
            'filename': filename,
            'type': ext,
            'size_kb': round(size / 1024, 2),
            'word_count': word_count,
            'keywords': ", ".join(keywords)
        })

    top_keywords = extract_keywords(all_text, 10)
    return summary, file_data, total_size, top_keywords

def save_report_txt(filepath, summary, file_data, total_size, top_keywords):
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write("Folder Summary Report\n\n")
        f.write(f"Total Files: {sum(summary.values())}\n")
        f.write(f"Total Size: {round(total_size / (1024 * 1024), 2)} MB\n\n")
        f.write("File Type Breakdown:\n")
        for k, v in summary.items():
            f.write(f"- {k.upper()}: {v}\n")
        f.write("\nTop Keywords:\n")
        for kw in top_keywords:
            f.write(f"- {kw}\n")
        f.write("\nFile Details:\n")
        for file in file_data:
            f.write(f"{file['filename']} ({file['type']}) - {file['size_kb']} KB, {file['word_count']} words, Keywords: {file['keywords']}\n")

def save_report_docx(filepath, summary, file_data, total_size, top_keywords):
    doc = DocxDocument()
    doc.add_heading("Folder Summary Report", level=1)
    doc.add_paragraph(f"Total Files: {sum(summary.values())}")
    doc.add_paragraph(f"Total Size: {round(total_size / (1024 * 1024), 2)} MB")
    doc.add_heading("File Type Breakdown", level=2)
    for k, v in summary.items():
        doc.add_paragraph(f"- {k.upper()}: {v}")
    doc.add_heading("Top Keywords", level=2)
    for kw in top_keywords:
        doc.add_paragraph(f"- {kw}")
    doc.add_heading("File Details", level=2)
    for file in file_data:
        doc.add_paragraph(f"{file['filename']} ({file['type']}) - {file['size_kb']} KB, {file['word_count']} words, Keywords: {file['keywords']}")
    doc.save(filepath)

if __name__ == "__main__":
    root = Tk()
    root.withdraw()

    folder_path = filedialog.askdirectory(title="Select Folder to Summarize")
    if not folder_path:
        messagebox.showinfo("Cancelled", "No folder selected. Exiting.")
        exit()

    output_path = filedialog.asksaveasfilename(title="Save Report As", defaultextension=".docx",
                                               filetypes=[("Word Document", "*.docx"), ("Text File", "*.txt")])
    if not output_path:
        messagebox.showinfo("Cancelled", "No output path selected. Exiting.")
        exit()

    summary, file_data, total_size, top_keywords = summarize_folder(folder_path)

    if output_path.endswith(".txt"):
        save_report_txt(output_path, summary, file_data, total_size, top_keywords)
    else:
        save_report_docx(output_path, summary, file_data, total_size, top_keywords)

    messagebox.showinfo("Done", f"Summary saved to {output_path}")
